import json
from pathlib import Path

import pandas as pd
from bs4 import BeautifulSoup
from openpyxl import load_workbook

from .config import SUPPORTED_EXTENSIONS
from .core import tokenize_text

try:
    from docx import Document
except ImportError:
    Document = None


TEXT_EXTENSIONS = {"txt", "log", "md", "conf", "cfg", "ini", "css"}
CELL_DELIMITERS = ["|", "\t", "，", ",", ";", "；"]


def get_correct_extension(filename, supported_extensions=SUPPORTED_EXTENSIONS):
    sorted_exts = sorted(supported_extensions, key=lambda x: -len(x))
    for ext in sorted_exts:
        if ext.lower() in filename.lower():
            return ext.lower()
    return None


def read_file_rows(file_path, extension):
    if extension in TEXT_EXTENSIONS:
        return read_text_rows(file_path)
    if extension == "csv":
        return read_csv_rows(file_path)
    if extension in {"xls", "xlsx"}:
        return read_excel_rows(file_path)
    if extension == "docx":
        return read_docx_rows(file_path)
    if extension in {"html", "htm", "xml"}:
        return read_html_xml_rows(file_path)
    if extension == "json":
        return read_json_rows(file_path)
    raise ValueError(f"不支持的文件类型: {extension}")


def iter_file_row_chunks(file_path, extension, chunk_size):
    chunk = []
    for row in iter_file_rows(file_path, extension):
        chunk.append(row)
        if len(chunk) >= chunk_size:
            yield chunk
            chunk = []
    if chunk:
        yield chunk


def iter_file_rows(file_path, extension):
    if extension in TEXT_EXTENSIONS:
        yield from iter_text_rows(file_path)
    elif extension == "csv":
        yield from iter_csv_rows(file_path)
    elif extension == "xlsx":
        yield from iter_xlsx_rows(file_path)
    else:
        yield from read_file_rows(file_path, extension)


def iter_text_rows(file_path):
    last_error = None
    for encoding in ["utf-8", "gbk", "gb2312", "latin-1"]:
        try:
            with open(file_path, "r", encoding=encoding, errors="replace") as file_obj:
                for line in file_obj:
                    text = line.strip()
                    row_data = normalize_row_values([text])
                    if len(row_data) == 1 and row_data[0] == text:
                        row_data = tokenize_text(text)
                    if row_data:
                        yield row_data
            return
        except Exception as exc:
            last_error = exc
    raise last_error


def iter_csv_rows(file_path):
    for chunk in pd.read_csv(file_path, low_memory=False, header=None, chunksize=5000):
        for _, row in chunk.iterrows():
            values = [str(cell).strip() if pd.notna(cell) else "" for cell in row]
            row_data = normalize_row_values(values)
            if any(row_data):
                yield row_data


def iter_xlsx_rows(file_path):
    workbook = load_workbook(file_path, read_only=True, data_only=True)
    try:
        for sheet in workbook.worksheets:
            for row in sheet.iter_rows(values_only=True):
                values = [str(cell).strip() if cell is not None else "" for cell in row]
                row_data = normalize_row_values(values)
                if any(row_data):
                    yield row_data
    finally:
        workbook.close()


def read_text_rows(file_path):
    return list(iter_text_rows(file_path))


def split_cell_value(value):
    value = str(value).strip()
    if not value:
        return []
    for delimiter in CELL_DELIMITERS:
        if delimiter in value:
            parts = [part.strip() for part in value.split(delimiter)]
            return [part for part in parts if part]
    return [value]


def normalize_row_values(values):
    row_data = []
    for value in values:
        row_data.extend(split_cell_value(value))
    return row_data


def read_csv_rows(file_path):
    data = []
    df = pd.read_csv(file_path, low_memory=False, header=None)
    if df.empty:
        return data
    for _, row in df.iterrows():
        values = [str(cell).strip() if pd.notna(cell) else "" for cell in row]
        row_data = normalize_row_values(values)
        if any(row_data):
            data.append(row_data)
    return data


def read_excel_rows(file_path):
    data = []
    suffix = Path(file_path).suffix.lower()
    engine = "openpyxl" if suffix == ".xlsx" else "xlrd"
    df = pd.read_excel(file_path, engine=engine, header=None)
    if df.empty:
        return data
    for _, row in df.iterrows():
        values = [str(cell).strip() if pd.notna(cell) else "" for cell in row]
        row_data = normalize_row_values(values)
        if any(row_data):
            data.append(row_data)
    return data


def read_docx_rows(file_path):
    if Document is None:
        raise RuntimeError("缺少 python-docx 依赖，请先安装 python-docx")
    data = []
    doc = Document(file_path)
    for para in doc.paragraphs:
        tokens = tokenize_text(para.text.strip())
        if tokens:
            data.append(tokens)
    for table in doc.tables:
        for row in table.rows:
            row_cells = normalize_row_values(cell.text.strip() for cell in row.cells)
            if any(row_cells):
                data.append(row_cells)
    return data


def read_html_xml_rows(file_path):
    try:
        with open(file_path, "r", encoding="utf-8", errors="replace") as file_obj:
            content = file_obj.read()
        soup = BeautifulSoup(content, "html.parser")
        data = []
        for line in soup.get_text(separator=" ").splitlines():
            tokens = tokenize_text(line.strip())
            if tokens:
                data.append(tokens)
        return data
    except Exception:
        return read_text_rows(file_path)


def read_json_rows(file_path):
    def extract_values(obj, results):
        if isinstance(obj, dict):
            for value in obj.values():
                extract_values(value, results)
        elif isinstance(obj, list):
            for item in obj:
                extract_values(item, results)
        else:
            value = str(obj).strip()
            if value:
                results.append(value)

    try:
        with open(file_path, "r", encoding="utf-8", errors="replace") as file_obj:
            json_obj = json.load(file_obj)
    except json.JSONDecodeError:
        return read_text_rows(file_path)

    data = []
    if isinstance(json_obj, list):
        for item in json_obj:
            row_values = []
            extract_values(item, row_values)
            if row_values:
                data.append(row_values)
    else:
        row_values = []
        extract_values(json_obj, row_values)
        if row_values:
            data.append(row_values)
    return data
